import argparse
import logging
import os
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Cm, RGBColor, Mm
from docx.shared import Pt

parser = argparse.ArgumentParser(
    prog="Создатель Приложений А и Б",
    description="Создаёт документ в формате ворда",
    epilog="- Foxelyss",
)
parser.add_argument("--application", action='store',default="Б",nargs='?', help="Буква приложения")
parser.add_argument("--file", action='store',default="application_b.docx",nargs='?', help="Файл записи")
args, unknownargs = parser.parse_known_args()

application_letter = args.application
file_name = args.file
logger = logging.getLogger(__name__)

logging.basicConfig(level=logging.INFO)

document = Document()

section = document.sections[0]

section.page_width = Mm(210)  # A4 210 mm
section.page_height = Mm(297)  # A4 297 mm

section.left_margin = Mm(30)    # левое – 30 мм
section.right_margin = Mm(15)   # правое – 15 мм
section.top_margin = Mm(20)     # верхнее – 20 мм
section.bottom_margin = Mm(20)  # нижнее – 20 мм

logger.info("Пишем в приложение "+ application_letter)
document.add_heading("ПРИЛОЖЕНИЕ " + application_letter, 1)

styles = document.styles

styles["Default Paragraph Font"].font.name = 'PT Astra Serif'
styles["Normal"].font.name = 'PT Astra Serif'
styles["Normal"].font.size = Pt(14)
styles["Normal"].paragraph_format.first_line_indent = Cm(1.25)
styles["Normal"].paragraph_format.space_after = Cm(0)
styles["Normal"].paragraph_format.space_before = Cm(0)
styles["Normal"].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

styles["Header"].font.name = 'PT Astra Serif'
styles["Header"].font.size = Pt(14)
styles["Header"].paragraph_format.first_line_indent = Cm(1.25)
styles["Header"].paragraph_format.space_after = Cm(0)
styles["Header"].paragraph_format.space_before = Cm(0)
styles["Header"].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

styles["Title"].font.name = 'PT Astra Serif'
styles["Title"].font.size = Pt(14)
styles["Title"].paragraph_format.first_line_indent = Cm(1.25)
styles["Title"].paragraph_format.space_after = Cm(0)
styles["Title"].paragraph_format.space_before = Cm(0)
styles["Title"].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


styles["Heading 1"].font.name = None
styles["Heading 1"].font.name = 'PT Astra Serif'
styles["Heading 1"].font.bold = False
styles["Heading 1"].font.size = Pt(14)
styles["Heading 1"].font.color.rgb = RGBColor(0, 0, 0)
styles["Heading 1"].paragraph_format.space_after = Cm(0)
styles["Heading 1"].paragraph_format.space_before = Cm(0)
styles["Heading 1"].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
styles["Heading 1"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

styles.add_style("code", WD_STYLE_TYPE.PARAGRAPH)

styles["code"].font.name = 'PT Astra Serif'
styles["code"].font.size = Pt(12)
styles["code"].paragraph_format.first_line_indent = Cm(0)#Cm(1.25)
styles["code"].paragraph_format.space_after = Cm(0)
styles["code"].paragraph_format.space_before = Cm(0)
styles["code"].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

max_length = 70
max_lines = 47

line_count = 0
for x in unknownargs:

    if os.path.isdir(x) or not os.path.isfile(x):
        logger.warning("Не файл: " + x)
        continue

    logger.info("Читаю файл: " + x)

    document.add_paragraph("", style='code')
    line_count += 1

    if line_count > max_lines:
        document.add_page_break()
        a = document.add_paragraph("ПРОДОЛЖЕНИЕ ПРИЛОЖЕНИЯ " + application_letter)
        a.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        line_count = 0

    document.add_paragraph(x[x.find("code")+4:], style='code')
    line_count += 1

    if line_count > max_lines:
        document.add_page_break()
        a = document.add_paragraph("ПРОДОЛЖЕНИЕ ПРИЛОЖЕНИЯ " + application_letter)
        a.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        line_count = 0

    document.add_paragraph("", style='code')
    line_count += 1

    if line_count > max_lines:
        document.add_page_break()
        a = document.add_paragraph("ПРОДОЛЖЕНИЕ ПРИЛОЖЕНИЯ " + application_letter)
        a.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        line_count = 0

    with open(x, "r") as file:
        for i in file.readlines():
            line_count += 1 + len(i.replace(" ","")) // max_length
            document.add_paragraph(i.replace("\n",""), style='code')

            if line_count > max_lines:
                document.add_page_break()
                a= document.add_paragraph("ПРОДОЛЖЕНИЕ ПРИЛОЖЕНИЯ "+ application_letter)
                a.alignment=WD_ALIGN_PARAGRAPH.RIGHT

                line_count = 0

logger.info("Файл создан: "+ file_name)

document.save(file_name)
